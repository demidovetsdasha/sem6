import os
import spacy
import docx2txt
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from tkinter import ttk
from docx import Document
from dataclasses import dataclass, asdict

# Загрузите модель spaCy для английского языка
nlp = spacy.load("en_core_web_sm")

@dataclass
class TokenData:
    text: str
    lemma: str
    pos: str
    tag: str
    morph: dict
    dep: str
    head: str

def extract_token_info(text):
    """
    Обрабатывает английский текст и возвращает список объектов TokenData,
    для каждого из которых указана:
      - оригинальная форма,
      - лемма,
      - универсальный тег части речи,
      - детальный тег,
      - морфологические признаки,
      - синтаксическая зависимость,
      - текст-голова (Head)
    Результат сортируется по алфавиту (без учета регистра).
    """
    doc = nlp(text)
    tokens_info = []
    for token in doc:
        # Оставляем только слова (без цифр, знаков препинания и пробелов)
        if token.is_alpha:
            token_obj = TokenData(
                text=token.text,
                lemma=token.lemma_,
                pos=token.pos_,
                tag=token.tag_,
                morph=token.morph.to_dict(),
                dep=token.dep_,
                head=token.head.text
            )
            tokens_info.append(token_obj)
    tokens_info_sorted = sorted(tokens_info, key=lambda x: x.text.lower())
    return tokens_info_sorted

def read_text_from_file(file_path):
    """
    Извлекает текст из файла DOC или DOCX.
    Для DOCX используется docx2txt.
    Для DOC используется textract.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        try:
            text = docx2txt.process(file_path)
            return text
        except Exception as e:
            messagebox.showerror("Ошибка при чтении DOCX", str(e))
            return None
    elif ext == ".doc":
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            return text
        except ImportError:
            messagebox.showerror("Ошибка", "Для чтения DOC файлов установите пакет textract.")
            return None
        except Exception as e:
            messagebox.showerror("Ошибка при чтении DOC", str(e))
            return None
    else:
        messagebox.showerror("Ошибка", "Неподдерживаемый формат файла. Выберите файл DOC или DOCX.")
        return None

def write_output_to_docx(output_path, tokens_info):
    """
    Создает DOCX документ и записывает в него информацию по каждой лексеме.
    """
    doc = Document()
    doc.add_heading("Анализ текста", level=0)
    for token in tokens_info:
        line = (f"Текст: {token.text}, Лемма: {token.lemma}, "
                f"POS: {token.pos}, Tag: {token.tag}, Morph: {token.morph}, "
                f"Dep: {token.dep}, Head: {token.head}")
        doc.add_paragraph(line)
    try:
        doc.save(output_path)
    except Exception as e:
        messagebox.showerror("Ошибка при сохранении", str(e))

class TokenApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Анализ текста и лексический словарь")
        self.geometry("900x600")
        
        self.tokens = []  # основной список лексем
        self.filtered_tokens = []  # список после фильтрации

        self.create_menu()
        self.create_widgets()

    def create_menu(self):
        menubar = tk.Menu(self)
        # Меню "Файл"
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Открыть файл", command=self.load_file)
        file_menu.add_command(label="Сохранить результат", command=self.save_result)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self.quit)
        menubar.add_cascade(label="Файл", menu=file_menu)
        
        # Меню "Редактирование"
        edit_menu = tk.Menu(menubar, tearoff=0)
        edit_menu.add_command(label="Добавить лексему", command=self.add_token)
        menubar.add_cascade(label="Редактирование", menu=edit_menu)
        
        # Меню "Помощь"
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="О программе", command=self.show_help)
        menubar.add_cascade(label="Помощь", menu=help_menu)
        
        self.config(menu=menubar)

    def create_widgets(self):
        # Фрейм для поисковой строки
        search_frame = tk.Frame(self)
        search_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(search_frame, text="Поиск:").pack(side=tk.LEFT)
        self.search_entry = tk.Entry(search_frame)
        self.search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.search_entry.bind("<KeyRelease>", lambda event: self.filter_tokens())

        # Кнопка для обновления фильтра
        tk.Button(search_frame, text="Сбросить фильтр", command=self.reset_filter).pack(side=tk.RIGHT)

        # Фрейм для Treeview
        tree_frame = tk.Frame(self)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        columns = ("text", "lemma", "pos", "tag", "morph", "dep", "head")
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings")
        for col in columns:
            self.tree.heading(col, text=col.capitalize())
            self.tree.column(col, width=100, anchor="center")
        self.tree.pack(fill=tk.BOTH, expand=True)

        # Двойной клик для редактирования выбранной лексемы
        self.tree.bind("<Double-1>", self.edit_token)

    def load_file(self):
        messagebox.showinfo("Выбор файла", "Выберите входной файл (DOC или DOCX)")
        input_file = filedialog.askopenfilename(
            title="Выберите входной файл",
            filetypes=[("DOC files", "*.doc"), ("DOCX files", "*.docx")]
        )
        if not input_file:
            messagebox.showerror("Ошибка", "Входной файл не выбран!")
            return

        text = read_text_from_file(input_file)
        if text is None:
            return

        # Обработка текста и получение списка лексем
        self.tokens = extract_token_info(text)
        self.filtered_tokens = self.tokens.copy()
        self.refresh_treeview()
        messagebox.showinfo("Успех", "Текст успешно обработан.")

    def refresh_treeview(self):
        # Очистка текущего содержимого
        for i in self.tree.get_children():
            self.tree.delete(i)
        # Заполнение данными
        for token in self.filtered_tokens:
            self.tree.insert("", tk.END, values=(
                token.text, token.lemma, token.pos, token.tag, str(token.morph), token.dep, token.head
            ))

    def filter_tokens(self):
        query = self.search_entry.get().lower()
        if not query:
            self.filtered_tokens = self.tokens.copy()
        else:
            self.filtered_tokens = [
                token for token in self.tokens
                if query in token.text.lower() or query in token.lemma.lower()
            ]
        self.refresh_treeview()

    def reset_filter(self):
        self.search_entry.delete(0, tk.END)
        self.filtered_tokens = self.tokens.copy()
        self.refresh_treeview()

    def save_result(self):
        messagebox.showinfo("Сохранение результата", "Выберите место для сохранения результата (DOCX)")
        output_file = filedialog.asksaveasfilename(
            title="Сохранить результат",
            defaultextension=".docx",
            filetypes=[("DOCX files", "*.docx")]
        )
        if not output_file:
            messagebox.showerror("Ошибка", "Файл для сохранения не выбран!")
            return

        write_output_to_docx(output_file, self.tokens)
        messagebox.showinfo("Успех", f"Результат успешно сохранён в {output_file}")

    def add_token(self):
        # Диалог для добавления новой лексемы
        new_data = self.token_dialog()
        if new_data:
            self.tokens.append(new_data)
            self.filtered_tokens = self.tokens.copy()
            self.refresh_treeview()

    def edit_token(self, event):
        # Редактирование выбранной лексемы по двойному клику
        item = self.tree.focus()
        if not item:
            return
        values = self.tree.item(item, "values")
        # Создаем объект TokenData на основе текущих значений
        current_token = TokenData(
            text=values[0],
            lemma=values[1],
            pos=values[2],
            tag=values[3],
            morph=eval(values[4]),  # осторожно, здесь используется eval для восстановления словаря
            dep=values[5],
            head=values[6]
        )
        new_data = self.token_dialog(current_token)
        if new_data:
            # Найти индекс изменяемой лексемы в списке
            index = self.tokens.index(current_token) if current_token in self.tokens else None
            if index is not None:
                self.tokens[index] = new_data
                self.filtered_tokens = self.tokens.copy()
                self.refresh_treeview()

    def token_dialog(self, token=None):
        """
        Диалоговое окно для ввода/редактирования данных лексемы.
        Если token передан, поля заполняются его значениями.
        Возвращает объект TokenData или None, если ввод отменен.
        """
        dialog = tk.Toplevel(self)
        dialog.title("Лексема")
        dialog.grab_set()  # делаем окно модальным

        labels = ["Текст", "Лемма", "POS", "Tag", "Morph (словарь)", "Dep", "Head"]
        entries = {}

        for idx, label in enumerate(labels):
            tk.Label(dialog, text=label + ":").grid(row=idx, column=0, sticky="e", padx=5, pady=5)
            entry = tk.Entry(dialog, width=50)
            entry.grid(row=idx, column=1, padx=5, pady=5)
            entries[label] = entry

        # Если передана лексема, заполняем поля
        if token:
            entries["Текст"].insert(0, token.text)
            entries["Лемма"].insert(0, token.lemma)
            entries["POS"].insert(0, token.pos)
            entries["Tag"].insert(0, token.tag)
            entries["Morph (словарь)"].insert(0, str(token.morph))
            entries["Dep"].insert(0, token.dep)
            entries["Head"].insert(0, token.head)

        def on_ok():
            try:
                # Преобразуем значение для поля Morph в словарь
                morph_val = eval(entries["Morph (словарь)"].get())
                new_token = TokenData(
                    text=entries["Текст"].get(),
                    lemma=entries["Лемма"].get(),
                    pos=entries["POS"].get(),
                    tag=entries["Tag"].get(),
                    morph=morph_val,
                    dep=entries["Dep"].get(),
                    head=entries["Head"].get()
                )
                dialog.new_token = new_token
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Ошибка", f"Некорректный ввод данных: {e}")

        tk.Button(dialog, text="ОК", command=on_ok).grid(row=len(labels), column=0, pady=10)
        tk.Button(dialog, text="Отмена", command=dialog.destroy).grid(row=len(labels), column=1, pady=10)

        self.wait_window(dialog)
        return getattr(dialog, "new_token", None)

    def show_help(self):
        help_text = (
            "Приложение выполняет анализ английского текста и формирование словаря лексем.\n\n"
            "Функциональность:\n"
            "1. Загрузить файл DOC или DOCX (Файл -> Открыть файл).\n"
            "2. Автоматический анализ текста с помощью spaCy и отображение результата в таблице.\n"
            "3. Поиск по лексемам по полям 'Текст' или 'Лемма'.\n"
            "4. Редактирование лексем: дважды кликните по строке для редактирования.\n"
            "5. Добавление новой лексемы (Редактирование -> Добавить лексему).\n"
            "6. Сохранение результата в DOCX (Файл -> Сохранить результат).\n"
        )
        messagebox.showinfo("О программе", help_text)

def main():
    app = TokenApp()
    app.mainloop()

if __name__ == "__main__":
    main()
