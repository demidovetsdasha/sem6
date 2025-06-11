import spacy
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from dataclasses import dataclass
import csv
from docx import Document  

nlp = spacy.load("en_core_web_sm")

@dataclass
class TokenData:
    text: str
    lemma: str
    pos: str
    morph: dict
    dep: str
    head: str
    count: int = 1  

def extract_token_info(text):
    """
    Обрабатывает английский текст и возвращает список объектов TokenData.
    Для каждого уникального слова (без учёта регистра) вычисляется количество вхождений.
    Результат сортируется по алфавиту (без учёта регистра).
    """
    doc = nlp(text)
    tokens_dict = {}
    for token in doc:
        if token.is_alpha:
            key = token.text.lower()
            if key not in tokens_dict:
                tokens_dict[key] = TokenData(
                    text=token.text,
                    lemma=token.lemma_,
                    pos=token.pos_,
                    morph=token.morph.to_dict(),
                    dep=token.dep_,
                    head=token.head.text,
                    count=1
                )
            else:
                tokens_dict[key].count += 1
    tokens_info_sorted = sorted(tokens_dict.values(), key=lambda x: x.text.lower())
    return tokens_info_sorted

def read_text_from_docx(file_path):
    """
    Считывает и возвращает текст из выбранного DOCX-файла.
    """
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        messagebox.showerror("Ошибка при чтении файла", str(e))
        return None

class TokenApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Анализ текста и лексический словарь")
        self.geometry("950x650")
        
        self.tokens = []          # основной список лексем
        self.filtered_tokens = [] # список после фильтрации

        self.create_menu()
        self.create_widgets()

    def create_menu(self):
        menubar = tk.Menu(self)
        # Меню "Файл"
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Открыть DOCX файл", command=self.load_file)
        file_menu.add_command(label="Сохранить результат (DOCX)", command=self.save_result)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self.quit)
        menubar.add_cascade(label="Файл", menu=file_menu)
        
        menubar.add_command(label="Добавить", command=self.add_token)
        menubar.add_command(label="О программе", command=self.show_help)
        self.config(menu=menubar)

    def create_widgets(self):
        # Фрейм для фильтрации по слову и количеству
        filter_frame = tk.Frame(self)
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="Слово:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.word_filter_entry = tk.Entry(filter_frame, width=20)
        self.word_filter_entry.grid(row=0, column=1, padx=5, pady=5)
        self.word_filter_entry.bind("<KeyRelease>", lambda event: self.auto_filter())

        tk.Label(filter_frame, text="Количество:").grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.count_filter_entry = tk.Entry(filter_frame, width=10)
        self.count_filter_entry.grid(row=0, column=3, padx=5, pady=5)
        self.count_filter_entry.bind("<KeyRelease>", lambda event: self.auto_filter())

        # Фрейм для отображения таблицы Treeview
        tree_frame = tk.Frame(self)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        columns = ("text", "lemma", "pos", "morph", "dep", "head", "count")
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings")
        headings = {
            "text": "Текст",
            "lemma": "Лемма",
            "pos": "POS",
            "morph": "Morph",
            "dep": "Dep",
            "head": "Head",
            "count": "Количество"
        }
        for col in columns:
            self.tree.heading(col, text=headings[col])
            self.tree.column(col, width=100, anchor="center")
        self.tree.pack(fill=tk.BOTH, expand=True)

        # Кнопка для сохранения выбранных строк в CSV
        self.save_csv_button = tk.Button(self, text="Сохранить выбранные строки в CSV", command=self.save_selected_to_csv)
        self.save_csv_button.pack(pady=10)

        self.tree.bind("<Double-1>", self.edit_token)
        self.tree.bind("<Button-3>", self.show_full_text)

    def load_file(self):
        """
        Считывает текст из выбранного DOCX-файла, обрабатывает его и выводит результат.
        """
        file_path = filedialog.askopenfilename(
            title="Выберите DOCX файл",
            filetypes=[("Word files", "*.docx"), ("Word 97-2003", "*.doc")]
        )

        if not file_path:
            messagebox.showerror("Ошибка", "Файл не выбран!")
            return

        text = read_text_from_docx(file_path)
        if text is None:
            return

        self.tokens = extract_token_info(text)
        self.filtered_tokens = self.tokens.copy()
        self.refresh_treeview()
        messagebox.showinfo("Успех", "Текст успешно обработан.")

    def refresh_treeview(self):
        """
        Обновляет отображение данных в таблице Treeview.
        """
        for item in self.tree.get_children():
            self.tree.delete(item)
        for token in self.filtered_tokens:
            self.tree.insert("", tk.END, values=(
                token.text, token.lemma, token.pos, str(token.morph),
                token.dep, token.head, token.count
            ))

    def auto_filter(self):
        """
        Фильтрует список лексем в зависимости от введённых данных.
        """
        word_query = self.word_filter_entry.get().strip().lower()
        count_query = self.count_filter_entry.get().strip()

        if not word_query and not count_query:
            self.filtered_tokens = self.tokens.copy()
        elif word_query and count_query.isdigit():
            count_value = int(count_query)
            self.filtered_tokens = [
                token for token in self.tokens
                if (word_query in token.text.lower() or word_query in token.lemma.lower())
                   and token.count == count_value
            ]
        elif word_query:
            self.filtered_tokens = [
                token for token in self.tokens
                if word_query in token.text.lower() or word_query in token.lemma.lower()
            ]
        elif count_query.isdigit():
            count_value = int(count_query)
            self.filtered_tokens = [token for token in self.tokens if token.count == count_value]
        else:
            self.filtered_tokens = self.tokens.copy()
        
        self.refresh_treeview()

    def save_result(self):
        """
        Сохраняет результат анализа в DOCX-файл.
        """
        output_file = filedialog.asksaveasfilename(
            title="Сохранить результат как DOCX",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")]
        )
        if not output_file:
            messagebox.showerror("Ошибка", "Файл для сохранения не выбран!")
            return

        try:
            doc = Document()
            doc.add_heading("Анализ текста", 0)

            for token in self.tokens:
                line = (f"Текст: {token.text}, Лемма: {token.lemma}, "
                        f"POS: {token.pos}, Morph: {token.morph}, "
                        f"Dep: {token.dep}, Head: {token.head}, Количество: {token.count}")
                doc.add_paragraph(line)
            doc.save(output_file)
            messagebox.showinfo("Успех", f"Результат успешно сохранён в {output_file}")
        except Exception as e:
            messagebox.showerror("Ошибка при сохранении", str(e))

    def save_selected_to_csv(self):
        """
        Сохраняет выделенные строки в формате CSV.
        """
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showerror("Ошибка", "Не выбраны строки для сохранения!")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv")]
        )
        if not file_path:
            messagebox.showerror("Ошибка", "Файл для сохранения не выбран!")
            return
        
        try:
            with open(file_path, mode="w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["Текст", "Лемма", "POS", "Morph", "Dep", "Head", "Количество"])
                for item in selected_items:
                    values = self.tree.item(item)["values"]
                    writer.writerow(values)
            messagebox.showinfo("Успех", "Выбранные строки успешно сохранены в CSV.")
        except Exception as e:
            messagebox.showerror("Ошибка при сохранении", str(e))
            
    def add_token(self):
        """
        Открывает диалоговое окно для ввода новой лексемы.
        """
        new_data = self.token_dialog(token=None)
        if new_data:
            self.tokens.append(new_data)
            self.filtered_tokens = self.tokens.copy()
            self.refresh_treeview()

    def edit_token(self, event):
        """
        Открывает диалоговое окно для редактирования или удаления выбранной лексемы по двойному клику.
        """
        item = self.tree.focus()
        if not item:
            return
        values = self.tree.item(item, "values")
        try:
            current_token = TokenData(
                text=values[0],
                lemma=values[1],
                pos=values[2],
                morph=eval(values[3]),
                dep=values[4],
                head=values[5],
                count=int(values[6])
            )
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать объект лексемы: {e}")
            return
        result = self.token_dialog(token=current_token)
        if result:
            if result == "delete":
                self.delete_token(current_token)
            else:
                try:
                    index = self.tokens.index(current_token)
                except ValueError:
                    index = None
                if index is not None:
                    self.tokens[index] = result
                    self.filtered_tokens = self.tokens.copy()
                    self.refresh_treeview()

    def delete_token(self, token_to_delete):
        """
        Удаляет токен из списка и обновляет таблицу.
        """
        self.tokens = [token for token in self.tokens if token != token_to_delete]
        self.filtered_tokens = self.tokens.copy()
        self.refresh_treeview()

    def token_dialog(self, token=None):
        """
        Диалоговое окно для ввода/редактирования данных лексемы.
        Возвращает TokenData, "delete" или None.
        """
        dialog = tk.Toplevel(self)
        dialog.title("Лексема")
        dialog.update_idletasks()
        dialog.wait_visibility()
        dialog.grab_set()

        labels = ["Текст", "Лемма", "POS", "Morph (словарь)", "Dep", "Head", "Количество"]
        entries = {}

        for idx, label in enumerate(labels):
            tk.Label(dialog, text=label + ":").grid(row=idx, column=0, sticky="e", padx=5, pady=5)
            entry = tk.Entry(dialog, width=50)
            entry.grid(row=idx, column=1, padx=5, pady=5)
            entries[label] = entry

        if token:
            entries["Текст"].insert(0, token.text)
            entries["Лемма"].insert(0, token.lemma)
            entries["POS"].insert(0, token.pos)
            entries["Morph (словарь)"].insert(0, str(token.morph))
            entries["Dep"].insert(0, token.dep)
            entries["Head"].insert(0, token.head)
            entries["Количество"].insert(0, str(token.count))

        def on_ok():
            try:
                morph_val = eval(entries["Morph (словарь)"].get())
                count_val = int(entries["Количество"].get())
                new_token = TokenData(
                    text=entries["Текст"].get(),
                    lemma=entries["Лемма"].get(),
                    pos=entries["POS"].get(),
                    morph=morph_val,
                    dep=entries["Dep"].get(),
                    head=entries["Head"].get(),
                    count=count_val
                )
                dialog.new_token = new_token
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Ошибка", f"Некорректный ввод данных: {e}")

        tk.Button(dialog, text="ОК", command=on_ok).grid(row=len(labels), column=0, pady=10)
        if token:
            def on_delete():
                dialog.new_token = "delete"
                dialog.destroy()
            tk.Button(dialog, text="Удалить", command=on_delete).grid(row=len(labels), column=1, pady=10)
        else:
            tk.Button(dialog, text="Отмена", command=dialog.destroy).grid(row=len(labels), column=1, pady=10)

        self.wait_window(dialog)
        return getattr(dialog, "new_token", None)

    def show_full_text(self, event):
        """
        Обработчик правого клика для показа полного текста ячейки.
        """
        region = self.tree.identify("region", event.x, event.y)
        if region == "cell":
            col = self.tree.identify_column(event.x)
            row_id = self.tree.identify_row(event.y)
            if row_id and col:
                values = self.tree.item(row_id, "values")
                try:
                    col_index = int(col.replace("#", "")) - 1
                    if 0 <= col_index < len(values):
                        full_text = values[col_index]
                        messagebox.showinfo("Полный текст", full_text)
                except ValueError:
                    pass

    def show_help(self):
        """
        Показывает справочную информацию об использовании приложения.
        """
        help_text = (
            "Приложение выполняет анализ английского текста и формирование словаря лексем.\n\n"
            "Функциональность:\n"
            "1. Открытие DOCX файла через меню Файл -> Открыть DOCX файл.\n"
            "2. Анализ текста с использованием spaCy с вычислением количества вхождений для каждой лексемы.\n"
            "3. Отображение результата в таблице с колонками: Текст, Лемма, POS, Morph, Dep, Head, Количество.\n"
            "4. Фильтрация происходит автоматически по введённым данным в поля \"Слово\" и \"Количество\".\n"
            "5. Сохранение результата происходит в формате DOCX.\n"
            "6. Для добавления новой лексемы используйте пункт меню \"Добавить\".\n"
            "7. Редактирование записи – двойной клик по соответствующей строке.\n"
            "8. Для просмотра полного текста ячейки выполните правый клик по интересующей колонке строки.\n"
        )
        messagebox.showinfo("О программе", help_text)

def main():
    app = TokenApp()
    app.mainloop()

if __name__ == "__main__":
    main()
